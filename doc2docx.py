# -*- encoding:utf8 -*-
import os
import docx
from win32com import client as wc #导入模块
from docx.shared import RGBColor,Pt

path = os.path.abspath('.')+"\\"
# print(path)
# path = 'D:\\word\\' #文件路径
# print(path)
# input('a')

files = []
for file in os.listdir(path):
    if file.endswith(".doc"):
        files.append(path+file)
# print(files)

# 文件转换doc2docx
word = wc.Dispatch("Word.Application")  # 打开word应用程序
for file in files:
    doc = word.Documents.Open(file)  # 打开word文件
    doc.SaveAs2("{}x".format(file), 12)  # 另存为后缀为".docx"的文件，其中参数12指docx文件
    doc.Close()  # 关闭原来word文件
word.Quit()

"""


files =[]
for file in os.listdir(path):
    if file.endswith(".docx"):
        files.append(path+file)
#print(files)




def set_run(run,font_size,bold,color):
    run.font.size =font_size
    run.bold=bold
    run.font.color.rgb=color


def info_update():
    '''此函数用于批量替换合同中需要替换的信息
    doc:文件
    old_info和new_info：原文字和需要替换的新文字
    '''
    #读取段落中的所有run，找到需替换的信息进行替换
    
    for file in files:
        doc = docx.Document(file)
    
        
        #先查找需要更换的非表单的内容，更改格式
        for i in range(len(doc.paragraphs)):
                for r in doc.paragraphs[i].runs:
                    font_size = r.font.size
                    bold = r.bold
                    color = r.font.color.rgb
                    rest = r.text.split(old_info)
                    r.text = ''
                    for text in rest[:-1]:
                        run = doc.paragraphs[i].add_run(text=text)
                        set_run(run,font_size,bold,color)
                        run = doc.paragraphs[i].add_run(old_info)
                        run.font.size = font_size
                        run.bold = bold
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    run = doc.paragraphs[i].add_run(rest[-1])
                    set_run(run,font_size,bold,color)
        
        for para in doc.paragraphs: #
            for r in para.runs:
                r.text = r.text.replace(old_info, new_info) #替换信息
        
        
              
                   
        #读取表格中的所有单元格，找到需替换的信息进行替换，更改格式,表格格式命令不准确需要调整
        for i in range(len(doc.tables)):
                for cell in doc.tables[i].rows:
                    font_size = cell.font.size
                    bold = cell.bold
                    color = cell.font.color.rgb
                    rest = cell.text.split(old_info)
                    cell.text = ''
                    for text in rest[:-1]:
                        run = doc.tables[i].add_run(text=text)
                        set_run(run,font_size,bold,color)
                        run = doc.tables[i].add_run(old_info)
                        run.font.size = font_size
                        run.bold = bold
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    run = doc.tables[i].add_run(rest[-1])
                    set_run(run,font_size,bold,color)            
                   
        #读取表格中的所有单元格，找到需替换的信息进行替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                     cell.text = cell.text.replace(old_info, new_info ) #替换信息   
        
        doc.save(path+"{}".format(file.split("/")[-1]))  #保存文件
        
def file_save(): #文件转移
    for file in files:
        doc = docx.Document(file)        
        doc.save(path+"替换结果/{}".format(file.split("/")[-1]))  #保存文件
        
if __name__ == '__main__':
    while True:
        choice = int(input('是否继续替换：\n1.Y\n2.N\n'))
        if choice == 1:
            old_info = input('查找：')
            new_info = input('替换为：')
            info_update()
            
        elif choice == 2:
            #file_save()
            break
    
    
    


"""
